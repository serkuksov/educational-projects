import datetime
from threading import Thread

from docx import Document
import requests
from bs4 import BeautifulSoup

HEADERS = {
    "Accept": "*/*",
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/100.0.4896.160 YaBrowser/22.5.1.985 Yowser/2.5 Safari/537.36",
}


def get_count_pages(autor):
    response = variable.get(autor, headers=HEADERS, allow_redirects=True)
    if response:
        soup = BeautifulSoup(response.text, 'lxml')
        count_poems = int(soup.find('h2').previous_sibling.previous_sibling.find('b').text)
        count_pages = count_poems // 50 + 1
        return count_pages


def get_list_poems(autor):
    count_pages = get_count_pages(autor)

    list_poems = []
    for count_page in range(count_pages):
        params = {'s': count_page * 50}
        response = variable.get(autor, headers=HEADERS, allow_redirects=True, params=params)
        if response:
            soup = BeautifulSoup(response.text, 'lxml')
            poems_html = soup.find('ul', type='square').find_all('a')
            for poem_html in poems_html:
                name = poem_html.text
                link = 'https://stihi.ru' + poem_html['href']
                list_poems.append([name, link])
        print(f'Получено {count_page} из {count_pages} ссылок')
    return list_poems


def get_text_poem(link_poem):
    response = variable.get(link_poem, headers=HEADERS, allow_redirects=True)
    if response:
        soup = BeautifulSoup(response.text, 'lxml')
        text_poem = soup.find('div', class_='text').text
        return text_poem


def get_poems_in_file(list_poems):
    doc = Document()
    i = 1
    for poem in list_poems:
        name = poem[0]
        link = poem[1]
        text_poem = get_text_poem(link)
        doc.add_heading(name)
        doc.add_paragraph().add_run(text_poem)
        print(f'Собрано{i} из {len(list_poems)} стихов')
        i += 1
    time = datetime.datetime.today().strftime('%Y-%m-%d %H.%M.%S')
    doc.save(f'Poem {time}.docx')


def get_all_poems(avtor):
    list_poems = get_list_poems(avtor)
    get_poems_in_file(list_poems)


if __name__ == '__main__':
    autor = 'https://stihi.ru/avtor/wanter'
    variable = requests.Session()
    get_all_poems(autor)
    print('test')
